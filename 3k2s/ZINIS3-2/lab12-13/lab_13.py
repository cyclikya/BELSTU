import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import RGBColor


LINE_MARKS = {
    "0": " ",
    "1": "  "
}

COLOR_MARKS = {
    "0": RGBColor(0, 0, 0),
    "1": RGBColor(1, 1, 1)
}


def text_to_bits(text):
    data = text.encode("utf-8")
    length = len(data)
    length_bits = format(length, "032b")
    data_bits = "".join(format(byte, "08b") for byte in data)
    return length_bits + data_bits


def bits_to_text(bits):
    if len(bits) < 32:
        return ""

    length = int(bits[:32], 2)
    data_bits = bits[32:32 + length * 8]

    byte_values = []
    for i in range(0, len(data_bits), 8):
        byte = data_bits[i:i + 8]
        if len(byte) == 8:
            byte_values.append(int(byte, 2))

    return bytes(byte_values).decode("utf-8", errors="replace")


def get_all_runs(doc):
    runs = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                runs.append(run)
    return runs


def encode_line_length(doc, bits):
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    if len(paragraphs) < len(bits):
        raise ValueError("В контейнере недостаточно абзацев для метода изменения длины строки")

    for i, bit in enumerate(bits):
        text = paragraphs[i].text.rstrip()

        if bit == "0":
            paragraphs[i].text = text + LINE_MARKS["0"]
        else:
            paragraphs[i].text = text + LINE_MARKS["1"]


def decode_line_length(doc, count):
    paragraphs = [p for p in doc.paragraphs if p.text]

    bits = ""
    for paragraph in paragraphs:
        if len(bits) >= count:
            break

        text = paragraph.text

        if text.endswith("  "):
            bits += "1"
        elif text.endswith(" "):
            bits += "0"

    return bits


def encode_color(doc, bits):
    runs = get_all_runs(doc)

    if len(runs) < len(bits):
        raise ValueError("В контейнере недостаточно фрагментов текста для метода цвета")

    for i, bit in enumerate(bits):
        runs[i].font.color.rgb = COLOR_MARKS[bit]


def decode_color(doc, count):
    runs = get_all_runs(doc)

    bits = ""
    for run in runs:
        if len(bits) >= count:
            break

        color = run.font.color.rgb

        if color == RGBColor(1, 1, 1):
            bits += "1"
        elif color == RGBColor(0, 0, 0):
            bits += "0"

    return bits


def encode_message(message, container_path, output_path):
    bits = text_to_bits(message)

    half = len(bits) // 2
    line_bits = bits[:half]
    color_bits = bits[half:]

    doc = Document(container_path)

    encode_line_length(doc, line_bits)
    encode_color(doc, color_bits)

    doc.save(output_path)


def decode_message(path):
    doc = Document(path)

    first_32_line_bits = decode_line_length(doc, 16)
    first_32_color_bits = decode_color(doc, 16)

    first_32_bits = first_32_line_bits + first_32_color_bits

    if len(first_32_bits) < 32:
        raise ValueError("Не удалось прочитать длину скрытого сообщения")

    message_length = int(first_32_bits, 2)
    total_bits = 32 + message_length * 8

    half = total_bits // 2
    line_count = half
    color_count = total_bits - half

    line_bits = decode_line_length(doc, line_count)
    color_bits = decode_color(doc, color_count)

    bits = line_bits + color_bits

    return bits_to_text(bits)


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("lab 13")
        self.root.geometry("520x360")
        self.root.resizable(False, False)
        self.root.configure(bg="#f3f4f6")

        self.container_path = ""
        self.output_path = ""

        self.build()

    def build(self):
        frame = tk.Frame(self.root, bg="#f3f4f6")
        frame.pack(padx=20, pady=20, fill="both", expand=True)

        self.text = tk.Text(
            frame,
            height=7,
            width=55,
            font=("Consolas", 10),
            bg="white",
            fg="#111827",
            relief="solid",
            bd=1
        )
        self.text.pack()
        self.text.insert("1.0", "violetta")

        tk.Button(
            frame,
            text="Выбрать контейнер",
            command=self.choose_container,
            width=24
        ).pack(pady=(14, 6))

        tk.Button(
            frame,
            text="Встроить сообщение",
            command=self.encode_action,
            width=24
        ).pack(pady=6)

        tk.Button(
            frame,
            text="Извлечь сообщение",
            command=self.decode_action,
            width=24
        ).pack(pady=6)

        self.info = tk.Label(
            frame,
            text="Файл не выбран",
            bg="#f3f4f6",
            fg="#374151",
            font=("Segoe UI", 9),
            wraplength=460,
            justify="left"
        )
        self.info.pack(pady=(14, 0), anchor="w")

    def choose_container(self):
        path = filedialog.askopenfilename(
            filetypes=[("Word document", "*.docx")]
        )

        if path:
            self.container_path = path
            self.info.config(text=f"Контейнер: {path}")

    def encode_action(self):
        if not self.container_path:
            messagebox.showerror("Ошибка", "Выберите файл-контейнер")
            return

        message = self.text.get("1.0", tk.END).strip()

        if not message:
            messagebox.showerror("Ошибка", "Введите сообщение")
            return

        output_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")]
        )

        if not output_path:
            return

        try:
            encode_message(message, self.container_path, output_path)
            self.output_path = output_path
            self.info.config(text=f"Готово: {output_path}")
            messagebox.showinfo("Готово", "Сообщение встроено")
        except Exception as error:
            messagebox.showerror("Ошибка", str(error))

    def decode_action(self):
        path = filedialog.askopenfilename(
            filetypes=[("Word document", "*.docx")]
        )

        if not path:
            return

        try:
            message = decode_message(path)
            self.text.delete("1.0", tk.END)
            self.text.insert("1.0", message)
            self.info.config(text=f"Извлечено из файла: {path}")
            messagebox.showinfo("Готово", "Сообщение извлечено")
        except Exception as error:
            messagebox.showerror("Ошибка", str(error))


if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()